#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import time
import logging
import threading
from datetime import datetime, timedelta
import re
import json
from queue import Queue
from pathlib import Path
import uuid

from flask import Flask, request, render_template, session, redirect, url_for, flash, jsonify, g, Response
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_session import Session

# === python-dotenv で .env を読み取る
from dotenv import load_dotenv

# === Supabase を使う場合 ===
from supabase import create_client, Client

# 音声関係
import speech_recognition as sr
from pydub import AudioSegment

# 文字コード自動判定
import chardet

# Wordファイル生成用 (WindowsのWord呼び出しではなく、python-docx を使用)
from docx import Document
from docx.shared import Pt

# === OpenAI関連 ===
import openai
from openai.error import APIError, RateLimitError
# from openai import APIError, RateLimitError

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)

############################################
# Flaskアプリの初期化
############################################
app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'your-secret-key-here')

# Flask-Sessionの設定
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=1)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True

Session(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.session_protection = "strong"

############################################
# 1) .env から環境変数を読み込み
############################################
load_dotenv()

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_ANON_KEY = os.environ.get("SUPABASE_ANON_KEY")

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    logging.warning("【警告】OPENAI_API_KEY が設定されていません。要約は動作しない可能性があります。")
else:
    logging.info("OpenAI APIキーが正常に読み込まれました。")

# OpenAI のキーをセット
openai.api_key = OPENAI_API_KEY

############################################
# 2) Supabase を使う場合の初期化
############################################
supabase: Client = None
if SUPABASE_URL and SUPABASE_ANON_KEY:
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
        logging.info("Supabase クライアントを初期化しました。")
    except Exception as e:
        logging.error(f"Supabase の初期化に失敗: {e}")
else:
    logging.info("Supabase URL または KEY が設定されていないため、Supabase は使用しません。")

############################################
# その他設定
############################################
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

############################################
# 3) ユーザーモデル & ログイン周り
############################################
class User(UserMixin):
    def __init__(self, user_id, email):
        self.id = user_id
        self.email = email

@login_manager.user_loader
def load_user(user_id):
    """セッションからユーザー情報を取得"""
    if not user_id:
        return None
    # まずセッションから
    if 'user_id' in session and str(session['user_id']) == str(user_id):
        return User(user_id, session.get('email'))
    # Supabase認証を利用する例（必要に応じて適宜調整）
    try:
        if supabase:
            response = supabase.auth.get_user()
            if response and hasattr(response, 'user'):
                user_data = response.user
                if str(user_data.id) == str(user_id):
                    return User(user_id, user_data.email)
    except Exception as e:
        logging.error(f"ユーザー情報取得エラー: {e}")
    return None

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('minutes_list'))

    if request.method == 'POST':
        email = request.form['username']
        password = request.form['password']

        if not supabase:
            flash("認証システムに接続できません。")
            return render_template('login.html')

        try:
            result = supabase.auth.sign_in_with_password({
                "email": email,
                "password": password
            })
            user_data = result.user
            user_id = user_data.id
            user_email = user_data.email

            user_obj = User(user_id, user_email)
            login_user(user_obj, remember=True)

            session.permanent = True
            session['user_id'] = user_id
            session['email'] = user_email
            session.modified = True

            next_page = request.args.get('next')
            if next_page and next_page.startswith('/'):
                return redirect(next_page)
            return redirect(url_for('minutes_list'))

        except Exception as e:
            logging.error(f"【ログインエラー】: {e}")
            flash("ログインできません。メールアドレスとパスワードを確認してください。")
            return render_template('login.html')

    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

############################################
# 4) 音声 → テキスト 変換
############################################
def convert_to_wav(input_path, output_path):
    """Pydub + ffmpeg で任意フォーマットを WAV へ変換"""
    try:
        audio = AudioSegment.from_file(input_path)
        audio.export(output_path, format="wav")
        return True
    except Exception as e:
        logging.error(f"WAV変換エラー: {e}")
        return False

def transcribe_long_audio(file_path, language="ja-JP"):
    """
    SpeechRecognition + Google で文字起こし。
    30 秒ごとに分割して認識する例。
    """
    r = sr.Recognizer()
    r.dynamic_energy_threshold = True
    r.energy_threshold = 300
    max_retries = 3
    chunk_duration = 30.0

    try:
        audio_segment = AudioSegment.from_wav(file_path)
        total_duration_sec = len(audio_segment) / 1000.0
        logging.info(f"音声ファイル 全体 {total_duration_sec:.1f}秒")

        full_text = []
        current_offset = 0.0

        while current_offset < total_duration_sec:
            remaining = total_duration_sec - current_offset
            this_chunk = min(chunk_duration, remaining)

            logging.info(f"オフセット {current_offset:.1f}s ~ {this_chunk:.1f}s を認識中")
            with sr.AudioFile(file_path) as source:
                audio_data = r.record(source, offset=current_offset, duration=this_chunk)

            recognized_chunk = ""
            for attempt in range(max_retries):
                try:
                    recognized_chunk = r.recognize_google(
                        audio_data,
                        language=language,
                        show_all=False
                    )
                    break
                except sr.RequestError as e:
                    logging.warning(f"RequestError: {e} / リトライ {attempt+1}/{max_retries}")
                    time.sleep(2)
                except sr.UnknownValueError:
                    logging.warning("UnknownValueError: 音声がうまく認識されませんでした")
                    time.sleep(2)

            full_text.append(recognized_chunk)
            current_offset += this_chunk

        result_text = " ".join(t.strip() for t in full_text if t.strip())
        logging.info(f"【音声認識完了】文字数: {len(result_text)}")
        return result_text

    except Exception as e:
        logging.error(f"音声認識中エラー: {e}")
        return ""

############################################
# 5) OpenAI で要約
############################################
def chunk_text(text, max_chars=3000):
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunks.append(text[start:end])
        start = end
    return chunks

def call_gpt_api(prompt_text, system_msg="", model="gpt-3.5-turbo", temperature=0.7, max_retries=5):
    """OpenAI ChatCompletion API で要約"""
    if not openai.api_key:
        logging.error("OpenAI APIキー未設定")
        return ""

    messages = []
    if system_msg:
        messages.append({"role": "system", "content": system_msg})
    messages.append({"role": "user", "content": prompt_text})

    for attempt in range(max_retries):
        try:
            response = openai.ChatCompletion.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=1500,
            )
            return response.choices[0].message.content.strip()

        except RateLimitError:
            wait_time = (attempt + 1) * 10
            logging.warning(f"レートリミット到達。{wait_time}秒待機:リトライ({attempt+1}/{max_retries})")
            time.sleep(wait_time)
        except APIError as e:
            wait_time = (attempt + 1) * 3
            logging.warning(f"APIエラー: {e}。{wait_time}秒待機:リトライ({attempt+1}/{max_retries})")
            time.sleep(wait_time)
        except Exception as e:
            logging.error(f"予期せぬエラー: {e}")
            if attempt < max_retries - 1:
                time.sleep(3)
                continue
            return ""

    logging.error(f"{max_retries}回リトライ失敗")
    return ""

def summarize_text_in_chunks(entire_text):
    chunks = chunk_text(entire_text)
    summaries = []
    total_chunks = len(chunks)

    additional_instructions = """
【議事録作成時の指示】
1. 話し言葉を箇条書きなどで整理し重複は削除
2. 過度な砕け表現を抑えつつ分かりやすく
3. 『英語で考えて日本語で返答』してください
4. 発言者の個人名は伏せても良い
"""

    for i, chunk in enumerate(chunks):
        # 進捗の SSE 送信（もし使うなら）
        if hasattr(g, 'progress_queue'):
            progress = int((i / total_chunks) * 98)
            g.progress_queue.put({
                'progress': progress,
                'current': i + 1,
                'total': total_chunks,
                'estimated_time': (total_chunks - i - 1) * 30
            })

        base_prompt = "以下のテキストを議事録風に要約してください。\n\n"
        prompt = f"{additional_instructions}\n\n{base_prompt}\n{chunk}"

        summary = call_gpt_api(
            prompt_text=prompt,
            system_msg=(
                "あなたはプロの議事録作成者です。"
                "話し言葉を整理・重複を省き、日本語できちんとまとめてください。"
            ),
            model="gpt-3.5-turbo",
            temperature=0.5
        )
        summaries.append(summary)

    # 最終的に100%送信
    if hasattr(g, 'progress_queue'):
        g.progress_queue.put({
            'progress': 100,
            'current': total_chunks,
            'total': total_chunks,
            'estimated_time': 0
        })

    return "\n".join(summaries)

############################################
# 6) python-docx を用いた Word保存
############################################
def save_minutes_to_word(minutes_text, meeting_name="未入力"):
    """
    python-docx を使ってサーバーサイドで docx を生成し、/uploads に保存する例。
    """
    if not minutes_text.strip():
        logging.info("要約が空なのでスキップ")
        return None

    # 安全なファイル名生成
    now_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_meeting_name = re.sub(r'[\\/*?:"<>|]', '_', meeting_name)
    filename = f"{safe_meeting_name}_{now_str}.docx"
    save_path = os.path.join(UPLOAD_FOLDER, filename)

    doc = Document()

    # 余白設定など
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)

    # フォント設定
    style = doc.styles['Normal']
    style.font.name = 'MS Gothic'
    style.font.size = Pt(10.5)

    # ヘッダー的に簡単な情報
    doc.add_paragraph(f"会議名: {meeting_name}")
    doc.add_paragraph(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")

    # 実際の議事録
    doc.add_paragraph(minutes_text)

    doc.save(save_path)
    logging.info(f"Wordファイルを保存: {save_path}")
    return save_path

############################################
# 7) 各種エンドポイント
############################################
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/check_auth')
def check_auth():
    try:
        if current_user.is_authenticated:
            return jsonify({'authenticated': True, 'user_email': current_user.email})
        return jsonify({'authenticated': False})
    except Exception as e:
        logging.error(f"認証チェックエラー: {e}")
        return jsonify({'authenticated': False})

@app.route('/minutes_list')
@login_required
def minutes_list():
    if not supabase:
        return "Supabase が初期化されていません。"
    try:
        response = supabase.table("minutes_log").select("*").execute()
        data = response.data
        return render_template("minutes_list.html", minutes_data=data)
    except Exception as e:
        logging.error(f"【minutes_list エラー】 {e}")
        return f"エラー: {str(e)}"

@app.route('/minutes_detail/<id>')
@login_required
def minutes_detail(id):
    if not supabase:
        return {"error": "Supabaseが初期化されていません。"}, 500
    try:
        response = supabase.table("minutes_log").select("*").eq("id", id).execute()
        if not response.data:
            return {"error": "指定された議事録が見つかりません。"}, 404
        return response.data[0]
    except Exception as e:
        logging.error(f"【minutes_detail エラー】 {e}")
        return {"error": str(e)}, 500

@app.route('/enter_meeting_name', methods=['GET', 'POST'])
def enter_meeting_name():
    if request.method == 'POST':
        meeting_name = request.form.get('meetingName', '未入力')
        session['meeting_name'] = meeting_name
        return redirect(url_for('result'))
    return render_template('enter_meeting_name.html')

def format_time(seconds):
    if seconds < 60:
        return f"{seconds:.1f}秒"
    elif seconds < 3600:
        return f"{seconds/60:.1f}分"
    else:
        return f"{seconds/3600:.1f}時間"

# グローバルスコープに進捗用のQueueを定義
progress_queues = {}

@app.route('/progress')
def progress():
    def generate():
        # ユーザーごとにユニークなキューを作成
        queue_id = str(uuid.uuid4())
        progress_queues[queue_id] = Queue()
        
        try:
            while True:
                progress_data = progress_queues[queue_id].get()
                if 'error' in progress_data:
                    yield f"data: {json.dumps({'error': True})}\n\n"
                    break
                yield f"data: {json.dumps(progress_data)}\n\n"
        except GeneratorExit:
            # クリーンアップ
            if queue_id in progress_queues:
                del progress_queues[queue_id]
        except Exception as e:
            logging.error(f"進捗送信エラー: {e}")
            if queue_id in progress_queues:
                del progress_queues[queue_id]

    return Response(generate(), mimetype='text/event-stream')

@app.route('/submit_audio', methods=['POST'])
@login_required
def submit_audio():
    try:
        # 進捗の更新（gオブジェクトは使わない）
        for queue in progress_queues.values():
            queue.put({'progress': 0})

        start_time = time.time()

        meeting_name = request.form.get('meetingName', '未入力')
        create_name = request.form.get('createName', '未入力')
        session['meeting_name'] = meeting_name
        session['create_name'] = create_name
        session['timestamp'] = datetime.now().strftime("%Y-%m-%d")

        audio_file = request.files.get('audioFile')
        if not audio_file:
            return "音声ファイルがありません。"

        stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        audio_path = os.path.join(UPLOAD_FOLDER, f"{stamp}_{audio_file.filename}")
        audio_file.save(audio_path)

        # WAVへ変換
        wav_path = os.path.join(UPLOAD_FOLDER, f"{stamp}.wav")
        success = convert_to_wav(audio_path, wav_path)
        if not success:
            return "音声ファイルの変換に失敗しました"

        # 文字起こし
        st_trans = time.time()
        text_result = transcribe_long_audio(wav_path, "ja-JP")
        ed_trans = time.time()

        # 要約
        st_sum = time.time()
        minutes_result = summarize_text_in_chunks(text_result)
        ed_sum = time.time()

        total_time = time.time() - start_time

        session['text_result'] = text_result
        session['minutes_result'] = minutes_result
        session['transcribe_time_str'] = format_time(ed_trans - st_trans)
        session['summary_time_str'] = format_time(ed_sum - st_sum)
        session['total_time_str'] = format_time(total_time)

        # Supabase保存
        try:
            supabase.table("minutes_log").insert({
                "text_content": text_result,
                "minutes_result": minutes_result,
                "meeting_date": session['timestamp'],
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "meeting_name": session['meeting_name'],
                "create_name": session['create_name']
            }).execute()
        except Exception as e:
            logging.warning(f"Supabaseへの書き込み失敗: {e}")

        # 進捗の更新
        for queue in progress_queues.values():
            queue.put({'progress': 100})

        return redirect(url_for('result'))

    except Exception as e:
        logging.error(f"【submit_audio エラー】 {e}")
        # エラー通知
        for queue in progress_queues.values():
            queue.put({'error': True})
        flash(f'エラー: {str(e)}')
        return redirect(url_for('index'))

@app.route('/submit_text', methods=['POST'])
@login_required
def submit_text():
    try:
        if hasattr(g, 'progress_queue'):
            g.progress_queue.put({'progress': 0})

        start_time = time.time()

        meeting_name = request.form.get('meetingName', '未入力')
        create_name = request.form.get('createName', '未入力')
        session['meeting_name'] = meeting_name
        session['create_name'] = create_name
        session['timestamp'] = datetime.now().strftime("%Y-%m-%d")

        text_file = request.files.get('textFile')
        if not text_file:
            return "テキストファイルがありません。"

        stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        text_path = os.path.join(UPLOAD_FOLDER, f"{stamp}_{text_file.filename}")
        text_file.save(text_path)

        # 文字コード判定
        with open(text_path, 'rb') as f:
            raw_data = f.read()
        detected = chardet.detect(raw_data)
        guessed_enc = detected["encoding"] or "utf-8"

        try:
            with open(text_path, 'r', encoding=guessed_enc, errors='replace') as f:
                text = f.read()
        except:
            with open(text_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()

        st_sum = time.time()
        minutes_result = summarize_text_in_chunks(text)
        ed_sum = time.time()

        total_time = time.time() - start_time

        session['text_result'] = text
        session['minutes_result'] = minutes_result
        session['transcribe_time_str'] = "N/A"
        session['summary_time_str'] = format_time(ed_sum - st_sum)
        session['total_time_str'] = format_time(total_time)

        # Supabase保存
        try:
            supabase.table("minutes_log").insert({
                "text_content": text,
                "minutes_result": minutes_result,
                "meeting_date": session['timestamp'],
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "meeting_name": session['meeting_name'],
                "create_name": session['create_name']
            }).execute()
        except Exception as e:
            logging.warning(f"Supabaseへの書き込み失敗: {e}")

        return redirect(url_for('result'))

    except Exception as e:
        logging.error(f"【submit_text エラー】 {e}")
        if hasattr(g, 'progress_queue'):
            g.progress_queue.put({'error': True})
        flash(f'エラー: {str(e)}')
        return redirect(url_for('index'))

@app.route('/result')
@login_required
def result():
    text_result = session.get('text_result', '')
    minutes_result = session.get('minutes_result', '')
    transcribe_time_str = session.get('transcribe_time_str', 'N/A')
    summary_time_str = session.get('summary_time_str', 'N/A')
    total_time_str = session.get('total_time_str', 'N/A')
    return render_template(
        'result.html',
        text_result=text_result,
        minutes_result=minutes_result,
        transcribe_time_str=transcribe_time_str,
        summary_time_str=summary_time_str,
        total_time_str=total_time_str
    )

@app.route('/clear_session')
def clear_session():
    session.clear()
    return redirect(url_for('index'))

############################################
# Word保存API (例)
############################################
@app.route('/custom_save', methods=['POST'])
def custom_save():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'データがありません'}), 400

        text = data.get('text', '').strip()
        meeting_name = data.get('meeting_name', '未入力').strip()
        if not text:
            return jsonify({'success': False, 'error': 'テキストが空です'}), 400

        # python-docxでファイルを作成
        file_path = save_minutes_to_word(text, meeting_name)
        if not file_path:
            return jsonify({'success': False, 'error': 'Word保存に失敗'}), 500

        return jsonify({'success': True, 'path': file_path, 'message': 'ファイルを保存しました'})
    except Exception as e:
        logging.error(f"Word保存エラー: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

############################################
# アップロードフォルダの古いファイル削除
############################################
def cleanup_uploads(max_age_hours=24):
    try:
        now = datetime.now()
        for filename in os.listdir(UPLOAD_FOLDER):
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            if os.path.isfile(file_path):
                age_hours = (now - datetime.fromtimestamp(os.path.getmtime(file_path))).total_seconds() / 3600
                if age_hours > max_age_hours:
                    os.remove(file_path)
                    logging.info(f"古いファイル削除: {filename}")
    except Exception as e:
        logging.error(f"クリーンアップエラー: {e}")

def init_app():
    with app.app_context():
        cleanup_uploads()

############################################
# Render等での起動
############################################
if __name__ == '__main__':
    logging.info("Flaskサーバ起動")
    init_app()

    # Render 等のPaaSでは host='0.0.0.0' + PORT環境変数を使う
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
